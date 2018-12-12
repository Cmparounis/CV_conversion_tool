from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
import argparse
import os

parser = argparse.ArgumentParser()
parser.add_argument("filename", help="name of input file")
parser.add_argument("filename2", help ="name of output file")

args = parser.parse_args()

wordDoc1 = Document(args.filename)
docname = args.filename2

def find_table(wordDoc):
    check1 = True
    while(check1):
        for table in wordDoc.tables:
            print([cell.text for cell in table.rows[0].cells])
            check = True
            while(check):
                input1 = input("Is this the work experience table? YES/NO \n")
                if input1 == 'YES':
                    return table
                    check1 == False
                elif input1 == 'NO':
                    check = False
                else:
                    print("Invalid input")

def read_jobs(table):
    dict = {}
    j = 0
    for row in table.rows:
        j += 1
        if j > 1:
            dict[j] = [cell.text for cell in row.cells]

    return dict

def switch_to_EBRD(dict):
    document = Document()

    for key, value in dict.items():
        run = document.add_paragraph().add_run()
        table = document.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        font.italic = True
        pos = Inches(1.74)
        table.rows[0].cells[0].text = "Date: from (month / year) to (month / year):"
        table.rows[0].cells[0].width = pos
        table.rows[1].cells[0].text = "Location:"
        table.rows[1].cells[0].width = pos
        table.rows[2].cells[0].text = "Company:"
        table.rows[2].cells[0].width = pos
        table.rows[3].cells[0].text = "Position:"
        table.rows[3].cells[0].width = pos
        table.rows[4].cells[0].text = "Description:"
        table.rows[4].cells[0].width = pos
        incr = 0
        for str in value:
            table.rows[incr].cells[1].text = str
            table.rows[incr].cells[1].width = Inches(3.96)
            incr += 1
        run.add_break()
        format_text(document, table)
    document.save(docname)

def format_text(document, table):
    styles = document.styles
    stylename = [style.name for style in styles ]
    if "EBRD_1" not in stylename and "EBRD_1" not in stylename:
        style = styles.add_style("EBRD_1", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        font.italic = True

        style2 = document.styles.add_style("EBRD_2", WD_STYLE_TYPE.PARAGRAPH)
        style2.base_style = styles['Normal']
        font = style2.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

    for row in table.rows:
        for paragraph in row.cells[0].paragraphs:
            paragraph.style = 'EBRD_1'
        for paragraph in row.cells[1].paragraphs:
            paragraph.style = 'EBRD_2'

switch_to_EBRD(read_jobs(find_table(wordDoc1)))
print("Document ready at "+ os.getcwd() + "\\" + docname)
